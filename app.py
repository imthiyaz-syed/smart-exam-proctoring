import csv
import json
from sklearn.preprocessing import LabelEncoder
from flask import Flask, render_template, request, flash, redirect, url_for, session,jsonify,send_file
import numpy as np
import mysql.connector
import cv2, os
import pandas as pd
import pickle
import smtplib
from google import genai
from google.genai import types 
import datetime
import time
import random
from dotenv import load_dotenv
from PIL import Image
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import threading
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SelectField
from wtforms.validators import DataRequired, Length
from werkzeug.security import check_password_hash, generate_password_hash  # For secure password handling
from wtforms import StringField, SubmitField
from wtforms.validators import DataRequired
from wtforms import StringField, DateField
import cv2
from flask import Response
import re
import tkinter as tk
from tkinter import messagebox
import cv2
from ultralytics import YOLO
from docx import Document
from PyPDF2 import PdfReader
import os
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from yolo_webcam import live_webcam_detection
from logic import *
from yolo_webcam import live_webcam_detection, proctoring_state

load_dotenv()

# Global status trackers initialization
val_data = "Completed"
head_status = "Straight"

mydb = mysql.connector.connect(host="localhost", user="root", port="3306", passwd="Imthiyaz@786", database="face_biometric")
cursor = mydb.cursor()

app = Flask(__name__)
app.config['SECRET_KEY'] = 'navya'
app.config["TRAINING_IMAGE_PATH"] = os.path.join(os.getcwd(), "TrainingImage")
app.config["TRAINED_MODEL_FOLDER_PATH"] = os.path.join(os.getcwd(), "Trained_Model")
app.config["TRAINED_MODEL_PATH"] = os.path.join(os.getcwd(), "Trained_Model" , "Trainner.yml")

if not os.path.exists(app.config["TRAINING_IMAGE_PATH"]):
    os.makedirs(app.config["TRAINING_IMAGE_PATH"])

if not os.path.exists(app.config["TRAINED_MODEL_FOLDER_PATH"]):
    os.makedirs(app.config["TRAINED_MODEL_FOLDER_PATH"])

MAX_WARNINGS = 3


def reset_exam_security_state():
    proctoring_state["warnings"] = 0
    proctoring_state["phone_detections"] = 0
    proctoring_state["screenshot_attempts"] = 0
    proctoring_state["violations_count"] = 0
    proctoring_state["exam_terminated"] = False
    proctoring_state["termination_reason"] = ""


def terminate_exam(reason):
    proctoring_state["exam_terminated"] = True
    proctoring_state["termination_reason"] = reason
    proctoring_state["status"] = reason
    proctoring_state["active"] = False
    session["exam_terminated"] = True
    session["exam_termination_reason"] = reason
    session["exam_in_progress"] = False


def build_security_status():
    warnings = int(proctoring_state.get("warnings", 0))
    exam_terminated = bool(proctoring_state.get("exam_terminated", False) or session.get("exam_terminated", False))

    return {
        "warnings": warnings,
        "max_warnings": MAX_WARNINGS,
        "phone_detections": int(proctoring_state.get("phone_detections", 0)),
        "screenshot_attempts": int(proctoring_state.get("screenshot_attempts", 0)),
        "violations_count": int(proctoring_state.get("violations_count", 0)),
        "remaining_warnings": max(0, MAX_WARNINGS - warnings),
        "exam_terminated": exam_terminated,
        "termination_reason": proctoring_state.get("termination_reason", session.get("exam_termination_reason", "")),
        "status": proctoring_state.get("status", "Idle"),
        "face_count": int(proctoring_state.get("face_count", 0)),
        "multiple_faces": bool(proctoring_state.get("multiple_faces", False)),
        "no_face": bool(proctoring_state.get("no_face", False)),
        "looking_away": bool(proctoring_state.get("looking_away", False)),
        "looking_side": bool(proctoring_state.get("looking_side", False)),
        "looking_down": bool(proctoring_state.get("looking_down", False)),
        "cheating_flag": bool(proctoring_state.get("cheating_flag", False)),
    }

class CreateExamForm(FlaskForm):
    exam_name = StringField('Exam Name', validators=[DataRequired()])  # Exam name field
    exam_date = DateField('Exam Date', format='%Y-%m-%d', validators=[DataRequired()])  # Exam date field


# Load the saved YOLO model from your local system
model_path = 'yolov8_saved_model.pt' 
model = YOLO(model_path)  # Load the YOLO model

# Variables to track the previous x-coordinate and height of the person's head
previous_x = None
previous_height = None
sender_address = 'cse.takeoff@gmail.com'
sender_pass = 'digkagfgyxcjltup'

def send_mail(subject,receiver_address, mail_content):
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = subject
    message.attach(MIMEText(mail_content, 'plain'))
    ses = smtplib.SMTP('smtp.gmail.com', 587)
    ses.starttls()
    ses.login(sender_address, sender_pass)
    text = message.as_string()
    ses.sendmail(sender_address, receiver_address, text)
    ses.quit()

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/portfolio_details')
def portfolio_details():
    return render_template('portfolio-details.html')


@app.route('/Add_student', methods=["GET", "POST"])
def Add_student():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        pwd = request.form['pwd']
        cpwd = request.form['cpwd']
        pno = request.form['pno']
        addr = request.form['addr']
        uname = request.form['uname']
        otp1 = random.randint(0000, 9999)
        ts = time.time()
        date = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
        
        # Check if the email already exists
        sql = "select * from user_registration"
        result = pd.read_sql_query(sql, mydb)
        email1 = result['email'].values
        if email in email1:
            flash("Email already existed", "warning")
            return render_template('add_student.html')
        
        if pwd == cpwd:
            # Initialize camera
            cam = cv2.VideoCapture(0, cv2.CAP_DSHOW)  # Use cv2.CAP_DSHOW for better compatibility on Windows
            
            # Check if camera is available
            if not cam.isOpened():
                flash("Error: Could not open camera.", "danger")
                return render_template('add_student.html')

            harcascadePath = r"haarcascade\haarcascade_frontalface_default.xml"
            detector = cv2.CascadeClassifier(harcascadePath)
            sampleNum = 0

            while True:
                ret, img = cam.read()
                if not ret:
                    flash("Error: Failed to capture image.", "danger")
                    return render_template('add_student.html')

                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                faces = detector.detectMultiScale(gray, 1.3, 5)

                for (x, y, w, h) in faces:
                    cv2.rectangle(img, (x, y), (x + w, y + h), (255, 0, 0), 2)
                    sampleNum += 1
                    cv2.imwrite(f"TrainingImage/{name}.{otp1}.{sampleNum}.jpg", gray[y:y + h, x:x + w])

                cv2.imshow('frame', img)

                if cv2.waitKey(100) & 0xFF == ord('q') or sampleNum > 350:
                    break

            cam.release()
            cv2.destroyAllWindows()

            # Save the student data to the database
            sql = "INSERT INTO user_registration (sid, name, email, uname, pwd, pno, addr, d1) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
            val = (otp1, name, email, uname, pwd, pno, addr, date)
            cursor.execute(sql, val)
            mydb.commit()
            flash("Successfully Registered", "warning")
            return render_template('add_student.html')

        else:
            flash("Password and Confirm Password do not match", "warning")
            return render_template('add_student.html')

    return render_template('add_student.html')

@app.route('/Add_faculty', methods=["GET", "POST"])
def Add_faculty():
    if request.method == 'POST':
        fullname = request.form['fullname']
        email = request.form['email']
        department = request.form['department']
        emp_id = request.form['emp_id']
        pwd = request.form['pwd']
        cpwd = request.form['cpwd']
        mobile = request.form['mobile']
        otp1 = random.randint(1000, 9999)  # random 4-digit OTP
        ts = time.time()
        date = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')

        # Validate mobile number (should be exactly 10 digits)
        if not re.match(r'^\d{10}$', mobile):
            flash("Mobile number must be exactly 10 digits", "warning")
            return render_template('add_faculty.html')

        # Validate username (only letters and spaces allowed)
        if not re.match(r'^[A-Za-z ]+$', fullname):
            flash("Username should only contain letters and spaces", "warning")
            return render_template('add_faculty.html')

        # Validate department (only letters and spaces allowed)
        if not re.match(r'^[A-Za-z ]+$', department):
            flash("Department should only contain letters and spaces", "warning")
            return render_template('add_faculty.html')

        # Check if the email already exists in the database
        sql = "SELECT * FROM faculty_registration"
        result = pd.read_sql_query(sql, mydb)
        email1 = result['email'].values
        
        if email in email1:
            flash("Email already existed", "warning")
            return render_template('add_faculty.html')

        # Check if password and confirm password match
        if pwd == cpwd:
            # Insert faculty details into the database
            sql = "INSERT INTO faculty_registration (otp, username, email, department, emp_id, pwd, mobile, date) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
            val = (otp1, fullname, email, department, emp_id, pwd, mobile, date)
            cursor.execute(sql, val)
            mydb.commit()

            # Send confirmation email with OTP
            msg = f'Dear {fullname},\n\nYour Faculty ID is: {emp_id}\nYour OTP is: {otp1}\n\nRegards,\nOnline Faculty Authentication System'

            send_mail("Faculty Registration", email, msg)

            flash("Faculty successfully registered", "success")
            return redirect('Add_faculty')

        else:
            flash("Password and confirm password do not match", "warning")
            return render_template('add_faculty.html')

    return render_template('add_faculty.html')


@app.route('/faculty/profile/<emp_id>', methods=['GET'])
def faculty_profile(emp_id):
    sql = "SELECT * from faculty_registration WHERE emp_id=%s"
    cursor.execute(sql, (emp_id,))
    faculty_data = cursor.fetchone()

    if faculty_data:
        return render_template('faculty_profile.html', faculty_data=faculty_data)
    else:
        flash("Faculty not found", "warning")
        return redirect(url_for('faculty_home'))



# Example Login Form using Flask-WTF
class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    user_type = SelectField('User Type', choices=[('1', 'Admin'), ('2', 'User')], validators=[DataRequired()])
# Admin login route
class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])


@app.route('/adminlogin', methods=["GET", "POST"])
def adminlogin():
    form = LoginForm()  # Initialize the form

    if form.validate_on_submit():
        username = form.username.data
        password = form.password.data
        # Perform your authentication logic here
        if username == 'admin' and password == 'admin':  # Hardcoded check for simplicity
            # flash("Welcome Admin", "success")
            return redirect(url_for('adminhome'))
        else:
            flash("Invalid Admin Credentials", "danger")

    return render_template('admin_login.html', form=form)  # Pass form object to the template

@app.route('/studentlogin', methods=['GET', 'POST'])
def studentlogin():
    form = LoginForm()  # Initialize the form

    if form.validate_on_submit():
        username = form.username.data
        password = form.password.data

        # Use parameterized queries to prevent SQL injection
        sql = "SELECT * FROM user_registration WHERE uname=%s AND pwd=%s"
        cursor.execute(sql, (username, password))
        results = cursor.fetchall()

        if not results:
            flash("Invalid Email / Password", "danger")
            return render_template('studentlogin.html', form=form)  # Always pass the form back

        # Set session variables upon successful login
        if len(results) > 0:
            recognizer = cv2.face.LBPHFaceRecognizer_create()  # cv2.createLBPHFaceRecognizer()
            recognizer.read(r"Trained_Model\Trainner.yml")
            harcascadePath = r"Haarcascade\haarcascade_frontalface_default.xml"
            faceCascade = cv2.CascadeClassifier(harcascadePath)
            global cam
            cam = cv2.VideoCapture(0)
            font = cv2.FONT_HERSHEY_SIMPLEX
            pkl_file = open('label_encoder.pkl', 'rb')
            le = pickle.load(pkl_file)
            pkl_file.close()
            global tt
            count = []
            flag = 0
            det = 0
            global val_data, global_stop
            global_stop = False
            while True:
                _, frame = cam.read()
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                faces = faceCascade.detectMultiScale(gray, 1.2, 5)
                for (x, y, w, h) in faces:
                    cv2.rectangle(frame, (x, y), (x + w, y + h), (225, 0, 0), 2)
                    Id, conf = recognizer.predict(gray[y:y + h, x:x + w])
                    print(conf)
                    if (conf > 38):
                        flag += 1
                        person_name = "Unknown"
                        if flag==20:
                            flash("Unknown detection", "danger")
                            return render_template("studentlogin.html", form=form)  # Pass form back here
                    else:
                        det+=1
                        tt = le.inverse_transform([Id])
                        person_name=str(tt[0])

                        if det == 10:
                            cam.release()
                            cv2.destroyAllWindows()
                            session['sid'] = results[0][1]  # Assuming ID is in the second column
                            session['name'] = results[0][2]  # Assuming Name is in the third column
                            session['email'] = results[0][3]  # Assuming Email is in the fourth column
                            flash(f"Welcome, {session['name']}", "success")  # Flash success message

                            return redirect(url_for('studenthome'))  # Redirect to the student home page  

                    cv2.putText(frame,str(person_name), (x, y + h),font, 1, (255, 255, 255), 2)
                
                cv2.imshow('im', frame)
                if (cv2.waitKey(1) == ord('q')):
                    break

    # Always pass form back to the template
    return render_template('studentlogin.html', form=form)


@app.route('/studenthome')
def studenthome():
    return render_template('studenthome.html')


@app.route('/logout')
def logout():
    reason = session.get("exam_termination_reason", "")
    reset_exam_security_state()
    proctoring_state["active"] = False
    session.clear()
    if reason:
        flash(f"Exam terminated due to cheating detection: {reason}", "danger")
    else:
        flash("Successfully logged out", "success")
    return redirect(url_for('studentlogin'))



@app.route('/faculty_login', methods=['GET', 'POST'])
def faculty_login():
    form = LoginForm()  # Initialize the form

    if form.validate_on_submit():  # Only proceed if the form is submitted and valid
        username = form.username.data
        password = form.password.data

        # Perform your authentication logic here (use parameterized queries to prevent SQL injection)
        sql = "SELECT * FROM faculty_registration WHERE username=%s"
        cursor.execute(sql, (username,))  # Use tuple for parameterized query
        results = cursor.fetchall()
        print("sql_information:", results)

        if not results:
            flash("Invalid Username / Password", "danger")
            return redirect(url_for('faculty_login'))

        # Assuming the password is stored in plain text in the database (not hashed)
        stored_password = results[0][5]  # Assuming the password is in the 6th column (index 5)

        # Compare the entered password with the stored plain-text password
        if stored_password == password:
            session['fid'] = results[0][0]  # Assuming faculty ID is in the first column (index 0)
            session['name'] = results[0][1]  # Assuming faculty name is in the second column (index 1)
            session['email'] = results[0][2]  # Assuming faculty email is in the fourth column (index 3)
            session['username'] = username  # Assuming faculty email is in the fourth column (index 3)
            flash("Welcome, " + results[0][1], "success")  # Display faculty name
            return redirect(url_for('faculty_home'))  # Redirect to faculty home page
        else:
            flash("Invalid Username / Password", "danger")
            return redirect(url_for('faculty_login'))

    return render_template('faculty_login.html', form=form)  # Return the login page if form is not submitted
            

@app.route('/studentlogin', methods=['GET','POST'])
def student():
    pass

# Admin Home Route
@app.route('/adminhome', methods=['GET','POST'])
def adminhome():
    # adminlogout = 'admin_logged_in' in session
    # flash("Welcome Admin", "success")
    return render_template('adminhome.html')

# Admin Home Route
@app.route('/faculty_home', methods=['GET','POST'])
def faculty_home():
    # adminlogout = 'admin_logged_in' in session
    # flash("Welcome Faculty", "success")
    return render_template('faculty_home.html')





@app.route('/adminlogout')
def adminlogout():
    # Logic for logging out the admin (e.g., clearing session or cookies)
    flash("Successfully logged out", "success")
    return redirect(url_for('adminlogin'))


@app.route('/view_registrations')
def view_registrations():
    sql = "select * from user_registration  "
    x = pd.read_sql_query(sql, mydb)
    x = x.drop(['pwd'], axis=1)

    return render_template('view_registrations.html', row_val=x.values.tolist())


@app.route('/training', methods=['POST', 'GET'])
def training():
    le = LabelEncoder()
    faces, Id = getImagesAndLabels("TrainingImage")
    Id = le.fit_transform(Id)
    output = open('label_encoder.pkl', 'wb')
    pickle.dump(le, output)
    output.close()
    recognizer = cv2.face.LBPHFaceRecognizer_create()
    recognizer.train(faces, np.array(Id))
    recognizer.save("Trained_Model/Trainner.yml")
    flash("Model Trained Successfully", "success")
    return render_template('adminhome.html')


def getImagesAndLabels(path):
    imagePaths = [os.path.join(path, f) for f in os.listdir(path)]
    faces = []
    Ids = []
    for imagePath in imagePaths:
        if imagePath.endswith(".jpg") is False:
            continue
        pilImage = Image.open(imagePath).convert('L')
        imageNp = np.array(pilImage, 'uint8')
        Id = str(os.path.split(imagePath)[-1].split(".")[0])
        faces.append(imageNp)
        Ids.append(Id)
    return faces, Ids


@app.route('/add_question', methods=["GET","POST"])
def add_question():
    return render_template('add_question.html')

@app.route('/qsnback', methods=["GET","POST"])
def qsnback():
    if request.method == 'POST':
        qsn = request.form['qsn']
        opt1 = request.form['opt1']
        opt2 = request.form['opt2']
        opt3 = request.form['opt3']
        opt4 = request.form['opt4']
        ans = request.form['ans']
        sub = request.form['sub']
        
        # For exam_paper table, we assume 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'hh', and 'i' are the columns.
        # You might need to adjust the data being inserted based on your specific requirement.
        # Here, I will assume you want to store the question and options in the 'b', 'c', 'd', 'e', 'f' columns.
        # Adjust this part based on your table's requirements.

        a = '1'  # Just a placeholder, update as necessary
        b = qsn  # The question text
        c = opt1  # Option 1
        d = opt2  # Option 2
        e = opt3  # Option 3
        f = opt4  # Option 4
        g = ans   # The correct answer (this is just an example; adapt based on your needs)

        # Insert into qsn_ans table
        sql_qsn_ans = "INSERT INTO qsn_ans(qsn,opt1,opt2,opt3,opt4,ans,username,subject) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)"
        val_qsn_ans = (qsn, opt1, opt2, opt3, opt4, ans, session['username'],sub)
        cursor.execute(sql_qsn_ans, val_qsn_ans)
        mydb.commit()

        flash("Question added", "info")
        return render_template('add_question.html')
    
    return render_template('add_question.html')


@app.route('/view_questions')
def view_questions():
    # email=session.get('email')
    sql = "select * from qsn_ans "
    x = pd.read_sql_query(sql, mydb)
    x = x.drop(['id'], axis=1)
    return render_template('viewqsn.html', row_val=x.values.tolist())

@app.route('/viewqsn_faculty')
def viewqsn_faculty():
    # email=session.get('email')
    sql = "select hh,i from exam_paper where username='"+session['username']+"' order by id desc limit 1"
    x = pd.read_sql_query(sql, mydb)
    return render_template('viewqsn_faculty.html', row_val=x.values.tolist())

@app.route('/view_papers/<exam>')
def view_papers(exam):
    sql = "select * from exam_paper where username='"+session['username']+"' and hh='"+exam+"' order by id desc"
    x = pd.read_sql_query(sql, mydb)
    x = x.drop(['id'], axis=1)
    return render_template('view_papers.html', row_val=x.values.tolist())






@app.route('/create_exam_back', methods=["POST", "GET"])
def create_exam_back():
    if request.method == 'POST':
        exam_name = request.form['exam_name']  
        exam_date = request.form['exam_date']
        sub = request.form['course_code']
        # Select random questions from the database
        s = "SELECT * FROM qsn_ans where username='"+session['username']+"' and subject='"+sub+"' ORDER BY RAND()"
        questions = pd.read_sql_query(s, mydb)

        print(len(questions))
        v = len(questions)

        # Insert the exam paper data into the database
        for i in range(v):
            row = questions.iloc[i]
            sql = """
            INSERT INTO exam_paper (a, b, c, d, e, f, g, hh, i, username) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (
                int(row['id']), clean_question_text(row['qsn']), clean_option_text(row['opt1']),
                clean_option_text(row['opt2']), clean_option_text(row['opt3']), clean_option_text(row['opt4']),
                extract_answer_letter(row['ans']), exam_name, exam_date, session['username']
            ))
        mydb.commit()
        flash("Exam paper created successfully!", "success")
        return redirect(url_for('create_exam_back'))  # Redirect to avoid form resubmission
    
    data_query = "select subject from qsn_ans where username = '"+session['username']+"' group by subject"
    data = pd.read_sql_query(data_query, mydb)
    return render_template('create_exam.html', data=data.values.tolist())



@app.route('/view_exam')
def view_exam():
    mydb.cursor().execute('set sql_mode=""')
    sql = "SELECT * FROM exam_paper WHERE i >= CURRENT_DATE() GROUP BY hh"
    x = pd.read_sql_query(sql, mydb)
    x = x.drop(['a', 'b', 'c', 'd', 'e', 'f', 'g'], axis=1)
    
    # Add an index to each row
    x['index'] = range(1, len(x) + 1)
    
    # Pass the data to the template
    return render_template('view_exam.html', row_val=x.to_dict(orient='records'))


def TrackImages():
    global val_data, global_stop, head_status, cam
    try:
        head, objects, cheating, status, face_count = live_webcam_detection()
        val_data = status
        head_status = head
    except Exception as e:
        print(f"Webcam detection error: {e}")
        val_data = "Error"
        head_status = "Unknown"

# Route for live webcam feed
def gen_frames():
    camera = cv2.VideoCapture(0)  # Use the default webcam
    while True:
        success, frame = camera.read()
        if not success:
            break
        else:
            # Encode the frame as JPEG
            _, buffer = cv2.imencode('.jpg', frame)
            frame = buffer.tobytes()
            yield (b'--frame\r\n'
                   b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')

@app.route('/video_feed')
def video_feed():
    return Response(gen_frames(), mimetype='multipart/x-mixed-replace; boundary=frame')


@app.route('/take_test/<s>/<s1>/<s2>')
def take_test(s=0, s1="", s2=""):
    sid=session.get("sid")
    s="select count(*) from results where sid='"+str(sid)+"' and ename='"+str(s1)+"' and edate='"+str(s2)+"' "
    y=pd.read_sql_query(s,mydb)
    count=y.values[0][0]
    
    if count==0:
        reset_exam_security_state()
        session["exam_in_progress"] = True
        session["exam_terminated"] = False
        session["exam_termination_reason"] = ""

        sql = "select * from exam_paper where hh='" + s1 + "' and i='" + s2 + "'"
        x = pd.read_sql_query(sql, mydb)
        questions = []
        for idx, row in enumerate(x.to_dict(orient='records'), 1):
            questions.append({
                "question": f"{idx}. {clean_question_text(row.get('b', ''))}",
                "opt1": clean_option_text(row.get('c', '')),
                "opt2": clean_option_text(row.get('d', '')),
                "opt3": clean_option_text(row.get('e', '')),
                "opt4": clean_option_text(row.get('f', '')),
                "answer": extract_answer_letter(row.get('g', '')),
            })

        row1 = len(questions)
        dd = row1
        global tt, cam

        t1 = threading.Thread(target=TrackImages)
        t1.start()
        return render_template("take_test.html", s=s, s1=s1, s2=s2, questions=questions, r1=row1)
    else:
        flash("You have already attempted the exam","warning")
        return render_template("view_exam.html")

@app.route('/textback', methods=['GET', 'POST'])
def textback():
    global val_data, head_status
    # Initialize camera at the start if needed
    if request.method == 'POST':
        if session.get("exam_terminated"):
            flash("Exam already terminated due to cheating detection.", "danger")
            return redirect(url_for('logout'))

        lenn = request.form['dpr']
        s1 = request.form['s1']
        s2 = request.form['s2']
        sid = session.get('sid')
        email = session.get('email')
        name = session.get('name')

        # Stop the webcam thread loop gracefully
        proctoring_state["exam_terminated"] = True

        # Check if cam is initialized before calling release()
        if 'cam' in globals() and cam is not None and cam.isOpened():
            cam.release()
            cv2.destroyAllWindows()

        # Safely pull the active states generated by the Tracking Thread
        current_val_data = globals().get('val_data', "Completed")
        current_head_status = globals().get('head_status', "Straight")

        final_results = []  # Initialize a list to store the final results

        for ss in range(0, int(lenn)):
            sss = "myans" + str(ss)
            ca = "currans" + str(ss)

            ca1 = request.form[ca]
            sss1 = request.form[sss]

            # Insert query using current_val_data and current_head_status
            sql = """
                INSERT INTO results(sid, sname, semail, ename, edate, ca, ua, status, head_status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (str(sid), str(name), str(email), str(s1), str(s2), str(ca1), str(sss1), current_val_data, current_head_status))
            mydb.commit()

            # Append each result to the final results list for blockchain insertion
            final_results.append({
                'sid': sid,
                'sname': name,
                'semail': email,
                'ename': s1,
                'edate': s2,
                'ca': ca1,
                'ua': sss1,
                'status': current_val_data,
                'head_status': current_head_status
            })

        # Insert into finalresults using current_val_data and current_head_status
        sql1 = f"""
            INSERT INTO finalresults(sid, semail, ename, edate, ca, ua, status, head_status)
            VALUES ('{str(sid)}', '{str(email)}', '{str(s1)}', '{str(s2)}',
                    (SELECT count(*) FROM results WHERE ca = ua AND sid = '{str(sid)}' AND ename = '{str(s1)}' AND edate = '{str(s2)}'),
                    (SELECT count(*) FROM results WHERE sid = '{str(sid)}' AND ename = '{str(s1)}' AND edate = '{str(s2)}'),
                    '{current_val_data}', '{current_head_status}')
        """
        cursor.execute(sql1)
        mydb.commit()

        # Add the final results to the blockchain using addData function
        response = addData({"final_results": final_results})
        print(response)
        # Provide feedback to the user
        session["exam_in_progress"] = False
        session.pop("exam_terminated", None)
        session.pop("exam_termination_reason", None)
        reset_exam_security_state()
        flash("Your answers submitted. Exam Completed! Results added to the blockchain.", "warning")

    return render_template("view_exam.html")


# @app.route('/textback', methods=['GET', 'POST'])
# def textback():
#     # Initialize camera at the start if needed
#     if request.method == 'POST':
#         lenn = request.form['dpr']
#         s1 = request.form['s1']
#         s2 = request.form['s2']
#         sid = session.get('sid')
#         email = session.get('email')
#         name = session.get('name')

#         # Check if cam is initialized before calling release()
#         if 'cam' in globals() and cam.isOpened():
#             cam.release()
#             cv2.destroyAllWindows()

#         # Define val_data and head_status
#         val_data = "completed"  # Example value for val_data
#         head_status = "pending"  # Define head_status based on your logic

#         for ss in range(0, int(lenn)):
#             sss = "myans" + str(ss)
#             ca = "currans" + str(ss)

#             ca1 = request.form[ca]
#             sss1 = request.form[sss]

#             # Insert query using val_data and head_status
#             sql = """
#                 INSERT INTO results(sid, sname, semail, ename, edate, ca, ua, status, head_status)
#                 VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
#             """
#             cursor.execute(sql, (str(sid), str(name), str(email), str(s1), str(s2), str(ca1), str(sss1), val_data, head_status))
#             mydb.commit()

#         # Insert into finalresults using val_data and head_status
#         sql1 = f"""
#             INSERT INTO finalresults(sid, semail, ename, edate, ca, ua, status, head_status)
#             VALUES ('{str(sid)}', '{str(email)}', '{str(s1)}', '{str(s2)}',
#                     (SELECT count(*) FROM results WHERE ca = ua AND sid = '{str(sid)}' AND ename = '{str(s1)}' AND edate = '{str(s2)}'),
#                     (SELECT count(*) FROM results WHERE sid = '{str(sid)}' AND ename = '{str(s1)}' AND edate = '{str(s2)}'),
#                     '{val_data}', '{head_status}')
#         """
#         cursor.execute(sql1)
#         mydb.commit()

#         flash("Your answers submitted. Exam Completed!", "warning")

#     return render_template("view_exam.html")

@app.route('/exam_results')
def exam_results():
    mydb.cursor().execute('set sql_mode=""')
    sql = "SELECT * FROM exam_paper GROUP BY hh order by id desc"
    x = pd.read_sql_query(sql, mydb)
    x = x.drop(['a', 'b', 'c', 'd', 'e', 'f', 'g'], axis=1)
    print(x.values.tolist())  # Debug output
    return render_template('exam_results.html', row_val=x.values.tolist())


@app.route('/exam_results_back')
def exam_results_back():
    s = request.args.get('s', default=0, type=int)
    s1 = request.args.get('s1', default="", type=str)
    s2 = request.args.get('s2', default="", type=str)
    
    if not s1 or not s2:
        return "Invalid parameters", 400
    
    sq = "SELECT * FROM finalresults WHERE ename=%s AND edate=%s"
    x = pd.read_sql_query(sq, mydb, params=(s1, s2))
    return render_template("exam_results1.html", row_val=x.values.tolist())



@app.route('/reply_mail/<s>/<s1>/<s2>/<s3>/<s4>/')
def reply_mail(s=0, s1="", s2="",s3="",s4=""):
    if s4=="Cell Phone":
        msg = 'Cell Phone detected in the exam hall. Please do not use any electronic devices during the exam.'
        t = 'Regards,'
        t1 = 'Online Student Examination System.'
        mail_content = f'Dear {s},\n\n{msg}\n\n{t},\n{t1}'
    else:
        msg = 'You are authenticated and you can view your results'
        t = 'Regards,'
        t1 = 'Online Student Examination System.'
        mail_content = f'Dear {s},\n\n{msg}\n\n{t},\n{t1}'

    send_mail("Exam Results", s2, mail_content)
    return render_template("exam_results1.html")

@app.route('/view_result')
def view_result():
    sid = session.get('sid')

    sq = "SELECT * FROM finalresults WHERE sid='" + str(sid) + "' AND status != 'Cell Phone'"
    x = pd.read_sql_query(sq, mydb)

    # Drop the unnecessary columns
    x = x.drop(['id', 'sid', 'semail', 'status'], axis=1)

    # Convert the DataFrame into a list of dictionaries
    data = x.to_dict(orient='records')

    return render_template("view_result.html", row_val=data)


@app.route('/get_security_status')
def get_security_status():
    if not session.get("exam_in_progress"):
        return jsonify(build_security_status())

    if proctoring_state.get("exam_terminated"):
        terminate_exam(proctoring_state.get("termination_reason", "Cheating detected"))
        return jsonify(build_security_status())

    violation_detected = False
    termination_reason = ""

    if proctoring_state.get("cheating_flag"):
        proctoring_state["phone_detections"] = int(proctoring_state.get("phone_detections", 0)) + 1
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        proctoring_state["warnings"] = MAX_WARNINGS
        violation_detected = True
        termination_reason = "Cell phone detected"
    elif proctoring_state.get("multiple_faces"):
        proctoring_state["warnings"] = max(int(proctoring_state.get("warnings", 0)), 1)
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        violation_detected = True
        termination_reason = "Multiple faces detected"
    elif proctoring_state.get("no_face"):
        proctoring_state["warnings"] = max(int(proctoring_state.get("warnings", 0)), 1)
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        violation_detected = True
        termination_reason = "No face detected"
    elif proctoring_state.get("looking_side"):
        proctoring_state["warnings"] = MAX_WARNINGS
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        violation_detected = True
        termination_reason = "Looking Side"
    elif proctoring_state.get("looking_down"):
        proctoring_state["warnings"] = MAX_WARNINGS
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        violation_detected = True
        termination_reason = "Looking Down"
    elif proctoring_state.get("looking_away"):
        proctoring_state["warnings"] = max(int(proctoring_state.get("warnings", 0)), 1)
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        violation_detected = True
        termination_reason = "Looking away frequently"

    if violation_detected:
        terminate_exam(termination_reason)

    return jsonify(build_security_status())


@app.route('/check_exam_status')
def check_exam_status():
    if session.get("exam_terminated") or proctoring_state.get("exam_terminated"):
        return jsonify({
            "status": "terminated",
            "reason": proctoring_state.get("termination_reason", session.get("exam_termination_reason", "Cheating detected")),
        })

    return jsonify({"status": "active"})


@app.route('/report_tab_switch', methods=['POST'])
def report_tab_switch():
    if session.get("exam_in_progress"):
        proctoring_state["warnings"] = int(proctoring_state.get("warnings", 0)) + 1
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1

        if int(proctoring_state.get("warnings", 0)) >= MAX_WARNINGS:
            terminate_exam("Tab switching detected")

    return jsonify(build_security_status())


@app.route('/report_screenshot_attempt', methods=['POST'])
def report_screenshot_attempt():
    if session.get("exam_in_progress"):
        proctoring_state["screenshot_attempts"] = int(proctoring_state.get("screenshot_attempts", 0)) + 1
        proctoring_state["violations_count"] = int(proctoring_state.get("violations_count", 0)) + 1
        proctoring_state["warnings"] = MAX_WARNINGS
        terminate_exam("Screenshot attempt detected")

    return jsonify(build_security_status())

import os
from google import genai

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client(api_key=GOOGLE_API_KEY)

GENAI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-1.5-flash",

]


def generate_with_retry(prompt, retries=3):
    last_error = None

    for model_name in GENAI_MODELS:
        for attempt in range(retries):
            try:
                return client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
            except Exception as e:
                last_error = e
                error_text = str(e)
                is_retryable = any(code in error_text for code in ["503", "500", "UNAVAILABLE", "RESOURCE_EXHAUSTED"])

                print(f"Gemini request failed for model {model_name} on attempt {attempt + 1}: {e}")

                if is_retryable and attempt < retries - 1:
                    time.sleep(2 ** attempt)
                    continue

                break

    raise last_error


def clean_question_text(question):
    if not question:
        return ""

    cleaned = str(question).strip()
    cleaned = re.sub(r"^\s*(question|q)\s*\d+\s*[:.)-]\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*\d+\s*[:.)-]\s*", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    return cleaned


def clean_option_text(option):
    if not option:
        return ""

    cleaned = str(option).strip()
    cleaned = re.sub(r"^\s*[A-Da-d]\s*[:.)-]\s*", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def extract_answer_letter(answer):
    if not answer:
        return ""

    match = re.search(r"\b([a-dA-D])\b", str(answer))
    if match:
        return match.group(1).lower()

    return str(answer).strip().lower()[:1]

def generate_mcqs(text, number, tone):
    prompt = f"""
    Text: {text}
    
    You are an expert MCQ maker. Given the above text, create a quiz of {number} multiple choice questions in {tone} tone.
    Format like the example:
    
    1: Question?
    a) Option 1
    b) Option 2
    c) Option 3
    d) Option 4
    Correct answer: x
    """
    
    try:
        response = generate_with_retry(prompt)
        print("Response from AI:")
        print(response)
        quiz = response.text

        # Parse generated quiz
        mcqs = quiz.split('\n\n')
        quiz_data = []
        answer_key = {}
        for mcq in mcqs:
            lines = mcq.split('\n')
            if len(lines) < 6:
                continue
            question = clean_question_text(lines[0])
            options = {
                "a": clean_option_text(lines[1]),
                "b": clean_option_text(lines[2]),
                "c": clean_option_text(lines[3]),
                "d": clean_option_text(lines[4]),
            }
            correct = extract_answer_letter(lines[5])

            if not question or any(not option for option in options.values()) or correct not in ["a", "b", "c", "d"]:
                continue
            numbered_question = f"{len(quiz_data) + 1}. {question}"
            quiz_data.append({"Question": numbered_question, "Options": options})
            answer_key[numbered_question] = correct

        if not quiz_data:
            return None, None

        return quiz_data, answer_key

    except Exception as e:
        print(f"Error during question generation: {e}")
        return None, None
    
# def parse_question(question_str):
#     # Make sure question_str is a string
#     if not isinstance(question_str, str):
#         print(f"Invalid input: {question_str}")
#         return None

#     # Regex to separate the question from the options
#     match = re.match(r"\*\*(\d+):\s*(.*?)\*\*\s*a\) (.*?)\nb\) (.*?)\nc\) (.*?)\nd\) (.*)", question_str, re.DOTALL)
    
#     if match:
#         question = match.group(2)  # Extract the question
#         opt1 = match.group(3)      # Option a
#         opt2 = match.group(4)      # Option b
#         opt3 = match.group(5)      # Option c
#         opt4 = match.group(6)      # Option d
#         return question, opt1, opt2, opt3, opt4
#     else:
#         print(f"Failed to parse: {question_str}")
#     return None  # If no match found
# Function to parse the dictionary-based quiz data

def parse_question_data(quiz_entry):
    # Extracting the question and options from the dictionary
    question = clean_question_text(quiz_entry['Question'])
    options = quiz_entry['Options']
    
    # Ensure options are correctly extracted, defaulting to empty string if missing
    opt1 = clean_option_text(options.get('a', ''))
    opt2 = clean_option_text(options.get('b', ''))
    opt3 = clean_option_text(options.get('c', ''))
    opt4 = clean_option_text(options.get('d', ''))
    
    # Truncate question if it exceeds the max length
    max_question_length = 255
    if len(question) > max_question_length:
        question = question[:max_question_length]
    
    return question, opt1, opt2, opt3, opt4

@app.route('/prediction', methods=['GET', 'POST'])
def prediction():
    if request.method == "POST":
        try:
            text = None
            sub = request.form.get("sub", "").strip()
            txt = request.form.get("txt", "").strip()
            number_of_questions = int(request.form.get("number_of_questions", 5))
            tone = request.form.get("tone", "standard")
            
            print(f"Subject: {sub}")
            print(f"Number of questions: {number_of_questions}")
            print(f"Tone: {tone}")
            
            # Validate subject
            if not sub:
                flash("Please enter a subject.", "warning")
                return redirect(url_for("prediction"))
            
            # Handle file upload
            if "file" in request.files:
                file = request.files["file"]
                if file and file.filename != '':
                    try:
                        filename = file.filename.lower()
                        print(f"Processing file: {filename}")
                        
                        if filename.endswith('.txt'):
                            # Try different encodings for text files
                            try:
                                text = file.read().decode('utf-8')
                            except UnicodeDecodeError:
                                try:
                                    file.seek(0)  # Reset file pointer
                                    text = file.read().decode('latin-1')
                                except:
                                    file.seek(0)
                                    text = file.read().decode('cp1252')
                                    
                        elif filename.endswith('.pdf'):
                            from PyPDF2 import PdfReader
                            pdf_reader = PdfReader(file)
                            text = ""
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text
                                    
                        elif filename.endswith('.docx'):
                            from docx import Document
                            doc = Document(file)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            
                        else:
                            flash("Unsupported file format. Please upload .txt, .pdf, or .docx files.", "warning")
                            return redirect(url_for("prediction"))
                        
                        print(f"File loaded, text length: {len(text)}")
                        
                        if not text or len(text.strip()) == 0:
                            flash("The file appears to be empty or contains no readable text.", "warning")
                            return redirect(url_for("prediction"))
                            
                    except Exception as e:
                        print(f"File reading error: {str(e)}")
                        flash(f"Error reading file: {str(e)}", "danger")
                        return redirect(url_for("prediction"))
            
            # Use text from textarea if no file uploaded
            if not text and txt:
                text = txt
                print(f"Using textarea content, length: {len(text)}")
            
            if not text:
                flash("Please provide text either by uploading a file or entering text.", "warning")
                return redirect(url_for("prediction"))
            
            # Check if text is too short
            if len(text) < 50:
                flash("Please provide more text (at least 50 characters) to generate meaningful MCQs.", "warning")
                return redirect(url_for("prediction"))
            
            # Generate MCQs
            print("Calling generate_mcqs function...")
            quiz_data, answer_key = generate_mcqs(text, number_of_questions, tone)
            
            if quiz_data and len(quiz_data) > 0:
                print(f"Successfully generated {len(quiz_data)} MCQs")
                session['quiz_data'] = quiz_data
                session['answer_key'] = answer_key
                
                # Store questions in database
                success_count = 0
                for quiz_entry in quiz_data:
                    parsed_data = parse_question_data(quiz_entry)
                    
                    if parsed_data:
                        question, opt1, opt2, opt3, opt4 = parsed_data
                        answer = answer_key.get(quiz_entry['Question'], "")
                        if answer:
                            answer = answer.split(' ')[-1].strip().lower()
                            # Ensure answer is a, b, c, or d
                            if answer not in ['a', 'b', 'c', 'd']:
                                # Try to extract just the letter
                                if len(answer) > 0:
                                    answer = answer[0]
                                else:
                                    answer = 'a'
                        
                        try:
                            cursor.execute(''' 
                                INSERT INTO qsn_ans (qsn, opt1, opt2, opt3, opt4, ans, username, subject) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            ''', (question, opt1, opt2, opt3, opt4, answer, session.get('username', 'faculty'), sub))
                            mydb.commit()
                            success_count += 1
                        except Exception as e:
                            print(f"Database insert error: {e}")
                            mydb.rollback()
                
                flash(f"Successfully generated and saved {success_count} MCQs!", "success")
                return render_template("add_question2.html", quiz_data=quiz_data)
            else:
                print("Failed to generate MCQs - no data returned")
                flash("Failed to generate MCQs. Please check your API key and try again with different text.", "danger")
                return redirect(url_for("prediction"))
                
        except Exception as e:
            print(f"Prediction route error: {str(e)}")
            import traceback
            traceback.print_exc()
            flash(f"An error occurred: {str(e)}", "danger")
            return redirect(url_for("prediction"))
    
    return render_template("add_question2.html")

@app.route("/download/<format>")
def download(format):
    quiz_data = session.get("quiz_data", [])
    answer_key = session.get("answer_key", {})

    if format == "pdf":
        pdf_file = SimpleDocTemplate('mcqs.pdf', pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        for mcq in quiz_data:
            elements.append(Paragraph(mcq['Question'], styles['Heading1']))
            for option, text in mcq['Options'].items():
                elements.append(Paragraph(f"{option}) {text}", styles['Normal']))
            elements.append(Paragraph("", styles['Normal']))

        pdf_file.build(elements)
        return send_file("mcqs.pdf", as_attachment=True)

    elif format == "word":
        document = Document()
        for mcq in quiz_data:
            document.add_heading(mcq['Question'], level=1)
            for option, text in mcq['Options'].items():
                document.add_paragraph(f"{option}) {text}")
            document.add_paragraph()

        document.save("mcqs.docx")
        return send_file("mcqs.docx", as_attachment=True)

    elif format == "answer_key":
        answer_key_doc = Document()
        answer_key_doc.add_heading("Answer Key", level=1)
        for question, answer in answer_key.items():
            answer_key_doc.add_paragraph(f"{question} - Correct answer: {answer}")

        answer_key_doc.save("answer_key.docx")
        return send_file("answer_key.docx", as_attachment=True)

    return redirect(url_for("results"))

if __name__ == '__main__':
    app.run(debug=True , host='0.0.0.0')
